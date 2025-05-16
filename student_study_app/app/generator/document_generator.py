import logging
import os
import json
from docx import Document
from app.utils.logger import logger, log_msg
from app.processor.openai_processor import OpenAIProcessor

class DocumentGenerator:
    def __init__(self, output_dir):
        self.output_dir = output_dir
        self.openai_proc = OpenAIProcessor()
        self.quiz_answers = []

    def generate_week_material(self, openai_proc, feeds, current_grade, week_num=1):
        log_msg(f"Generating document for week {week_num}", level=logging.INFO)
        subject_map = {}
        other_topics = []
        self.quiz_answers = []

        raw_results = openai_proc.generate_educational_content(feeds, current_grade)
#         raw_results = r"""
# [
#     {
#         "subject_name": "Math Concepts",
#         "teacher_name": "Ellen Stenzler Whitford",
#         "date": "2025-05-11",
#         "topics": [
#             {
#                 "topic_name": "Area",
#                 "section_1_is_table": false,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": "Area is the amount of space inside a shape. For a rectangle, you calculate the area by multiplying its length by its width. Understanding area helps solve real-world problems, like figuring out how much carpet you need for a room.",
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Calculating the area of a garden to determine how many plants can fit.\n2. Measuring the area of a wall to know how much paint is needed.\n3. Finding the area of a tablecloth to cover a dining table.\n4. Determining the area of a playground.\n5. Calculating the area of a parking lot to plan the number of cars it can hold.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "How do you calculate the area of a rectangle?", "answer": "Multiply the length by the width" },
#                     { "question": "If a room is 5 meters long and 4 meters wide, what is its area?", "answer": "20 square meters" },
#                     { "question": "What is the area of a square with sides of 3 meters?", "answer": "9 square meters" },
#                     { "question": "Why is knowing the area important?", "answer": "It helps in planning and using space efficiently" },
#                     { "question": "How can you find the area of a triangle?", "answer": "Multiply the base by the height and divide by 2" }
#                 ]
#             },
#             {
#                 "topic_name": "Perimeter",
#                 "section_1_is_table": false,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": "Perimeter is the total distance around a shape. To find the perimeter of a rectangle, add up the lengths of all its sides. Knowing perimeter is useful for tasks like fencing a yard.",
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Calculating the perimeter of a garden to buy enough fencing.\n2. Measuring the perimeter of a picture frame.\n3. Determining the perimeter of a sports field.\n4. Finding the perimeter of a swimming pool.\n5. Measuring the perimeter of a room to install baseboards.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "How do you find the perimeter of a rectangle?", "answer": "Add the lengths of all four sides" },
#                     { "question": "What is the perimeter of a square with sides of 5 meters?", "answer": "20 meters" },
#                     { "question": "If a rectangle has sides of 6 meters and 4 meters, what is its perimeter?", "answer": "20 meters" },
#                     { "question": "Why is knowing the perimeter important?", "answer": "It helps in calculating the boundary length for enclosing spaces" },
#                     { "question": "How can you find the perimeter of a triangle?", "answer": "Add the lengths of all three sides" }
#                 ]
#             },
#             {
#                 "topic_name": "Fractions",
#                 "section_1_is_table": false,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": "Fractions represent parts of a whole. Understanding fractions is important for dividing things into equal parts and comparing sizes. For example, 1/2 is the same as 2/4.",
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Splitting a pizza into equal slices.\n2. Measuring ingredients for a recipe.\n3. Dividing a candy bar among friends.\n4. Comparing the sizes of different pieces of pie.\n5. Understanding half-time in sports.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "What is 1/2 equivalent to?", "answer": "2/4" },
#                     { "question": "How can you represent 3/4 on a number line?", "answer": "Divide a line into 4 equal parts and mark 3 parts" },
#                     { "question": "What fraction of a pizza is left if you eat 3 out of 8 slices?", "answer": "5/8" },
#                     { "question": "Why are fractions useful?", "answer": "They help in dividing things into equal parts and making comparisons" },
#                     { "question": "How can you add 1/4 and 2/4?", "answer": "By adding the numerators: 1 + 2 = 3/4" }
#                 ]
#             }
#         ],
#         "is_educational": true
#     },
#     {
#         "subject_name": "Vocabulary Development",
#         "teacher_name": "Rebecca Anastos",
#         "date": "2025-05-08",
#         "topics": [
#             {
#                 "topic_name": "Current Vocabulary Words",
#                 "section_1_is_table": true,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": [
#                     { "name": "ideal", "meaning": "something that is perfect", "example": "This vacation is ideal for relaxing." },
#                     { "name": "advertise", "meaning": "to give information or announce something is for sale", "example": "They advertise toys on TV." },
#                     { "name": "secluded", "meaning": "a place that is quiet and out of sight", "example": "The cabin was in a secluded forest." },
#                     { "name": "queasy", "meaning": "to feel sick to your stomach", "example": "She felt queasy on the bus." },
#                     { "name": "gleam", "meaning": "to shine or reflect a glow of light", "example": "The car's surface had a bright gleam." },
#                     { "name": "harmony", "meaning": "working well with others, getting along", "example": "The choir sang in perfect harmony." },
#                     { "name": "loaf", "meaning": "time spent relaxing, doing nothing", "example": "I like to loaf around on Sundays." },
#                     { "name": "object", "meaning": "to disagree", "example": "He objected to the new rules." },
#                     { "name": "jovial", "meaning": "always laughing and in a good mood", "example": "He is jovial and friendly to everyone." },
#                     { "name": "hardy", "meaning": "strong, able to survive in tough conditions", "example": "The hardy plant survived the winter." },
#                     { "name": "abbreviated", "meaning": "to write something in a shortened way", "example": "The word 'street' is often abbreviated as 'St.'" },
#                     { "name": "tedious", "meaning": "boring or repetitive", "example": "Filing papers can be a tedious job." },
#                     { "name": "diminish", "meaning": "making smaller", "example": "The noise began to diminish as the train left." },
#                     { "name": "convince", "meaning": "to persuade someone to think the way you do", "example": "She convinced him to try the new restaurant." },
#                     { "name": "interview", "meaning": "to ask someone questions", "example": "The journalist conducted an interview with the author." }
#                 ],
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Describing a perfect day as 'ideal'.\n2. Seeing advertisements on YouTube or billboards.\n3. Visiting a quiet, secluded cabin.\n4. Feeling queasy after a roller coaster ride.\n5. Watching the gleam of a polished car.\n6. Singing in harmony with a choir.\n7. Loafing on the couch after school.\n8. Objecting to a new rule in class.\n9. Meeting a jovial friend at the park.\n10. Observing a hardy plant in a garden.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "What does 'secluded' mean?", "answer": "A place that is quiet and private" },
#                     { "question": "What does 'jovial' describe?", "answer": "Someone who is always cheerful" },
#                     { "question": "What happens when you 'advertise'?", "answer": "You promote or announce something" },
#                     { "question": "Use 'queasy' in a sentence.", "answer": "I felt queasy before the test" },
#                     { "question": "What is the meaning of 'loaf'?", "answer": "To relax or do nothing" }
#                 ]
#             }
#         ],
#         "is_educational": true
#     },
#     {
#         "subject_name": "Reading Comprehension",
#         "teacher_name": "Rebecca Anastos",
#         "date": "2025-05-09",
#         "topics": [
#             {
#                 "topic_name": "Reading Comprehension Practice",
#                 "section_1_is_table": false,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": "Reading comprehension is the ability to understand and interpret what you read. Practicing with different types of texts and questions can improve this skill. Using strategies like paraphrasing questions and taking notes can help.",
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Reading a story and summarizing it in your own words.\n2. Answering questions about a passage you read in class.\n3. Discussing a book with friends and sharing your thoughts.\n4. Using context clues to understand new words in a text.\n5. Practicing reading comprehension with online resources.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "What is reading comprehension?", "answer": "The ability to understand and interpret what you read" },
#                     { "question": "How can you improve reading comprehension?", "answer": "By practicing with different texts and questions" },
#                     { "question": "What is one strategy to help with reading comprehension?", "answer": "Paraphrasing questions and taking notes" },
#                     { "question": "Why is reading comprehension important?", "answer": "It helps in understanding and learning from texts" },
#                     { "question": "What can you do if you encounter a difficult word while reading?", "answer": "Use context clues to understand its meaning" }
#                 ]
#             }
#         ],
#         "is_educational": true
#     },
#     {
#         "subject_name": "Social Studies",
#         "teacher_name": "Ellen Stenzler Whitford",
#         "date": "2025-05-11",
#         "topics": [
#             {
#                 "topic_name": "Culture",
#                 "section_1_is_table": false,
#                 "section_1_name": "Reading Material",
#                 "section_1_content": "Culture includes the beliefs, traditions, food, clothing, music, and celebrations of a group of people. Understanding different cultures helps us appreciate diversity and learn about the world.",
#                 "section_2_is_table": false,
#                 "section_2_name": "Real World Examples",
#                 "section_2_content": "1. Celebrating different holidays from around the world.\n2. Trying foods from different cultures.\n3. Learning about traditional clothing from various countries.\n4. Listening to music from different cultures.\n5. Participating in cultural festivals and events.",
#                 "section_3_is_table": false,
#                 "section_3_name": "Quizzes",
#                 "section_3_content": [
#                     { "question": "What is culture?", "answer": "The beliefs, traditions, food, clothing, music, and celebrations of a group of people" },
#                     { "question": "Why is it important to learn about different cultures?", "answer": "It helps us appreciate diversity and learn about the world" },
#                     { "question": "Give an example of a cultural celebration.", "answer": "Chinese New Year" },
#                     { "question": "What is one way to experience another culture?", "answer": "Trying foods from different cultures" },
#                     { "question": "How can music help us understand culture?", "answer": "It reflects the traditions and values of a culture" }
#                 ]
#             }
#         ],
#         "is_educational": true
#     }
# ]
# """
        log_msg(f"Raw OpenAI results: {raw_results}", level=logging.INFO)

        # Normalize and zip with feeds
        results = []
        if isinstance(raw_results, str):
            try:
                parsed = json.loads(raw_results)
                if isinstance(parsed, list):
                    results = parsed
            except json.JSONDecodeError as e:
                print(e)
                results = []
        elif isinstance(raw_results, list):
            results = raw_results

        log_msg(f"Processing {len(results)} enriched results from OpenAI", level=logging.INFO)

        for enriched_json in results:
            if enriched_json.get("is_educational", False):
                subject = enriched_json.get("subject_name", "General")
                if subject not in subject_map:
                    subject_map[subject] = enriched_json
                elif enriched_json.get('topics'):
                    subject_map[subject].get('topics',[]).append(enriched_json.get('topics'))
            else:
                # No corresponding feed, so cannot add to other_topics
                pass

        # Create document after all grouping is complete
        doc = Document()
        doc.add_heading(f"{current_grade} Grade Weekly Study Material - Week {week_num}", 0)

        for subject, enriched_map in subject_map.items():
            self._add_subject_section(doc, subject, enriched_map)

        self._add_other_topics_section(doc, other_topics)
        self._add_notes_section(doc, feeds)
        self._add_quiz_answers_section(doc)

        filename = os.path.join(self.output_dir, f"week{week_num}_material.docx")
        os.makedirs(self.output_dir, exist_ok=True)
        doc.save(filename)
        log_msg(f"Saved material to {filename}", level=logging.INFO)
        return filename

    def _add_subject_section(self, doc, subject, enriched_map):
        log_msg(f"Adding subject section: {subject}", level=logging.INFO)
        doc.add_heading(subject, level=1)
        date_str = enriched_map.get("date", "Unknown")
        doc.add_paragraph(f"Date: {date_str}", style="List Bullet")
        teacher = enriched_map.get("teacher_name", "Class Teacher")
        self._add_teacher_heading(doc, teacher)
        for topic in enriched_map.get('topics'):
            self._add_topic_content(doc, topic, subject)

    def _add_topic_content(self, doc, topic, subject):
        topic_name = topic.get("topic_name", topic.get("subject_name", "Topic"))
        self._add_topic_header(doc, topic_name)
        for i in range(1, 4):
            section_name = topic.get(f"section_{i}_name", f"Section {i}")
            is_table = topic.get(f"section_{i}_is_table", False)
            content = topic.get(f"section_{i}_content", {})
            doc.add_heading(section_name, level=3)
            self._render_section_content(doc, subject, topic_name, section_name, is_table, content)

    def _render_section_content(self, doc, subject, topic, section_name, is_table, content):
        if is_table and isinstance(content, list):
            table = doc.add_table(rows=1, cols=3)
            # Set all borders of the table
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            def set_table_borders(table):
                tbl = table._tbl
                tblPr = tbl.tblPr
                tblBorders = OxmlElement('w:tblBorders')

                for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)

                tblPr.append(tblBorders)

            set_table_borders(table)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Meaning'
            hdr_cells[2].text = 'Example'
            for row in content:
                name = row.get("name", "")
                meaning = row.get("meaning", "")
                example = row.get("example", "")
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = meaning
                row_cells[2].text = example
        # Updated quizzes parsing logic for list format
        elif isinstance(content, list) and section_name.lower() == "quizzes":
            for q in content:
                question = q.get("question", "").strip()
                answer = q.get("answer", "").strip()
                if question:
                    doc.add_paragraph(f"{question}", style="List Bullet")
                    doc.add_paragraph("A: ", style="List Bullet")
                    self.quiz_answers.append((subject, topic, question, answer))
        elif isinstance(content, dict):
            if section_name.lower() == "quizzes":
                all_quizzes = []
                if "questions" in content:
                    all_quizzes = content["questions"]
                elif "current_grade" in content or "next_grade" in content:
                    all_quizzes = content.get("current_grade", []) + content.get("next_grade", [])
                for q in all_quizzes:
                    doc.add_paragraph(f"{q.get('question')}", style="List Bullet")
                    doc.add_paragraph("A: ", style="List Bullet")
                    self.quiz_answers.append((subject, topic, q.get("question"), q.get("answer", "")))
            else:
                for k, v in content.items():
                    if k.lower() == "example":
                        doc.add_paragraph(f"Example: {v}")
                    else:
                        doc.add_paragraph(f"{k}: {v}")
        # Fallback for quizzes in plain list format
        elif section_name.lower() == "quizzes" and isinstance(content, list):
            for q in content:
                question = q.get("question", "").strip()
                answer = q.get("answer", "").strip()
                if question:
                    doc.add_paragraph(f"{question}", style="List Bullet")
                    doc.add_paragraph("A: ", style="List Bullet")
                    self.quiz_answers.append((subject, topic, question, answer))
        elif isinstance(content, str):
            doc.add_paragraph(content)

    def _add_teacher_heading(self, doc, teacher_name):
        doc.add_heading(f"Teacher: {teacher_name}", level=2)

    def _add_topic_header(self, doc, topic):
        doc.add_heading(f"Topic: {topic}", level=2)

    def _add_section(self, doc, section_title, content):
        doc.add_heading(section_title, level=3)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style="List Bullet")
        elif isinstance(content, str):
            doc.add_paragraph(content)

    def _add_quiz_section(self, doc, quiz):
        if quiz:
            doc.add_heading("Quizzes", level=3)
            for i, q in enumerate(quiz, 1):
                doc.add_paragraph(f"{i}. {q.get("question")}", style="List Bullet")
                answer = q.get("answer")
                if "___" not in q.get("question"):
                    doc.add_paragraph("A: ", style="List Bullet")

    def _add_quiz_answers_section(self, doc):
        if not self.quiz_answers:
            return
        doc.add_page_break()
        doc.add_heading("Appendix: Answers", level=1)
        subject_map = {}
        seen = set()
        for subject, topic, question, answer in self.quiz_answers:
            key = (subject, topic, question)
            if key in seen:
                continue
            seen.add(key)
            if subject not in subject_map:
                subject_map[subject] = {}
            if topic not in subject_map[subject]:
                subject_map[subject][topic] = []
            subject_map[subject][topic].append((question, answer))

        for subject, topics in subject_map.items():
            doc.add_heading(subject, level=2)
            for topic, qa_pairs in topics.items():
                doc.add_heading(f"Topic: {topic}", level=3)
                for question, answer in qa_pairs:
                    doc.add_paragraph(f"Q: {question}", style="List Bullet")
                    doc.add_paragraph(f"A: {answer}", style="List Bullet")

    def _add_other_topics_section(self, doc, other_topics):
        if other_topics:
            log_msg(f"Adding Other Topics section with {len(other_topics)} entries", level=logging.INFO)
            doc.add_heading("Other Topics", level=1)
            for topic in other_topics:
                doc.add_paragraph(f"• {topic}", style="List Bullet")

    def _add_notes_section(self, doc, feeds):
        notes = []
        for feed in feeds:
            if "note" in feed:
                notes.append(feed["note"])

        if notes:
            log_msg(f"Adding Important Notes section with {len(notes)} notes", level=logging.INFO)
            doc.add_heading("Important Notes", level=1)
            for note in notes:
                doc.add_paragraph(f"• {note}", style="List Bullet")