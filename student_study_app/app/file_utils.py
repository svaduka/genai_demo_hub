import mimetypes
import json
from typing import Union
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def json_to_docx(input_json):
    # Parse the JSON string into a dictionary
    data = json.loads(input_json)
    
    # Create a new Document
    doc = Document()
    
    # Add a title
    doc.add_heading('Candidate Job Match Report', level=1)
    
    # Create a table with three columns
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Job Description'
    hdr_cells[2].text = 'Candidate'
    
    # Populate the table
    for key, value in data.items():
        if isinstance(value, dict):
            # Handle nested dictionaries
            sub_values = list(value.values())
            if len(sub_values) == 2:  # Expected two keys: 'Required' and 'Candidate'
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = sub_values[0]  # 'Required' value
                row_cells[2].text = sub_values[1]  # 'Candidate' value
            else:
                raise ValueError(f"Nested dictionary for '{key}' must have exactly two keys.")
        else:
            # Handle simple key-value pairs
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value
            row_cells[2].text = ''  # Leave the Candidate column empty for non-nested values
    docx_io = BytesIO()
    doc.save(docx_io)
    return docx_io


def json_to_pdf(input_json):
    """
    Parse the given JSON into a PDF table and return it as BytesIO.

    Args:
        input_json (str): The JSON data to parse as a string.

    Returns:
        BytesIO: A BytesIO object containing the generated PDF.
    """
    try:
        # Parse the input JSON
        json_data = json.loads(input_json)

        # Prepare table data
        table_data = [["Component", "Job Description", "Candidate"]]  # Header row
        styles = getSampleStyleSheet()
        normal_style = styles["Normal"]

        for key, value in json_data.items():
            if isinstance(value, dict):  # Handle nested dictionaries
                sub_values = list(value.values())
                if len(sub_values) == 2:  # Expected two keys: 'Required' and 'Candidate'
                    table_data.append([
                        Paragraph(key, normal_style),
                        Paragraph(sub_values[0], normal_style),
                        Paragraph(sub_values[1], normal_style)
                    ])
                else:
                    raise ValueError(f"Nested dictionary for '{key}' must have exactly two keys.")
            else:  # Handle simple key-value pairs
                table_data.append([
                    Paragraph(key, normal_style),
                    Paragraph(value, normal_style),
                    ""
                ])  # Add a single value with an empty "Candidate" cell

        # Create a BytesIO buffer for the PDF
        pdf_buffer = BytesIO()

        # Set up the PDF document with landscape orientation
        pdf = SimpleDocTemplate(
            pdf_buffer, pagesize=landscape(letter),
            rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20
        )

        # Dynamically calculate column widths based on page width
        page_width = landscape(letter)[0] - 40  # Adjust for margins
        col_widths = [0.2 * page_width, 0.4 * page_width, 0.4 * page_width]  # Distribute width proportionally

        # Create the table
        table = Table(table_data, colWidths=col_widths, repeatRows=1)

        # Add style to the table
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ])
        table.setStyle(style)

        # Build the PDF
        pdf.build([table])

        # Seek to the beginning of the buffer
        pdf_buffer.seek(0)

        return pdf_buffer
    except Exception as e:
        raise ValueError(f"Error generating PDF table: {str(e)}")

def get_file_type(file: BytesIO, file_name: str) -> str:
    """
    Determine the type of an uploaded file based on its MIME type or extension.

    Args:
        file (BytesIO): The uploaded file object.
        file_name (str): The name of the uploaded file.

    Returns:
        str: The file type ("pdf", "docx", "text", or "unknown").
    """
    mime_type, _ = mimetypes.guess_type(file_name)
    file_extension = file_name.split(".")[-1].lower()

    if file_extension == "pdf":
        return "pdf"
    elif file_extension in ["docx", "doc"]:
        return "docx"
    elif file_extension in ["txt"]:
        return "text"
    elif mime_type and "text" in mime_type:
        return "text"
    else:
        return "unknown"


def extract_content_from_file(file: BytesIO, file_name: str) -> Union[str, None]:
    """
    Extract content from an uploaded file based on its type.

    Args:
        file (BytesIO): The uploaded file object.
        file_name (str): The name of the uploaded file.

    Returns:
        str: The extracted content from the file.
    """
    file_type = get_file_type(file, file_name)

    try:
        if file_type == "pdf":
            return extract_content_from_pdf(file)
        elif file_type == "docx":
            return extract_content_from_docx(file)
        elif file_type == "text":
            return extract_content_from_text(file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        raise ValueError(f"Error reading file '{file_name}': {str(e)}")


def extract_content_from_pdf(file: BytesIO) -> str:
    """
    Extract content from an uploaded PDF file.

    Args:
        file (BytesIO): The uploaded PDF file object.

    Returns:
        str: The extracted content from the PDF.
    """
    try:
        reader = PdfReader(file)
        content = ""
        for page in reader.pages:
            content += page.extract_text()
        return content.strip()
    except Exception as e:
        raise ValueError(f"Error extracting content from PDF: {str(e)}")


def extract_content_from_docx(file: BytesIO) -> str:
    """
    Extract content from an uploaded DOCX file.

    Args:
        file (BytesIO): The uploaded DOCX file object.

    Returns:
        str: The extracted content from the DOCX file.
    """
    try:
        document = Document(file)
        content = "\n".join([paragraph.text for paragraph in document.paragraphs])
        return content.strip()
    except Exception as e:
        raise ValueError(f"Error extracting content from DOCX: {str(e)}")


def extract_content_from_text(file: BytesIO) -> str:
    """
    Extract content from an uploaded text file.

    Args:
        file (BytesIO): The uploaded text file object.

    Returns:
        str: The extracted content from the text file.
    """
    try:
        return file.read().decode("utf-8").strip()
    except Exception as e:
        raise ValueError(f"Error extracting content from text file: {str(e)}")

def parse_json(input_data: Union[str, dict]) -> dict:
    """
    Parse the input to ensure it's a JSON object.

    Args:
        input_data (str or dict): The input data, either as a JSON string or dictionary.
    
    Returns:
        dict: The parsed JSON object.
    """
    if isinstance(input_data, str):
        return json.loads(input_data)
    elif isinstance(input_data, dict):
        return input_data
    else:
        raise ValueError("Input must be a JSON string or dictionary.")
        

def convert_to_html_table(json_data: dict) -> str:
    """
    Convert a JSON object into an HTML table.

    Args:
        json_data (dict): The JSON object to be converted.
    
    Returns:
        str: HTML string representing the table.
    """
    table_html = "<table border='1' style='border-collapse: collapse; width: 50%; text-align: left;'>"
    table_html += "<tr><th>Key</th><th>Value</th></tr>"

    for key, value in json_data.items():
        table_html += f"<tr><td>{key}</td><td>{value}</td></tr>"

    table_html += "</table>"
    return table_html


def parse_nested_json_to_html_table(json_data):
    """
    Parse the given JSON into an HTML table.

    Args:
        json_data (dict): The JSON data to parse.

    Returns:
        str: An HTML string representing the table.
    """
    try:
        # Start the HTML table
        html = "<table border='1' style='border-collapse: collapse; width: 50%; text-align: left;'>"
        html += "<tr><th>Component</th><th>Job Description</th><th>Candidate</th></tr>"
        # Iterate through the JSON data and build table rows
        for key, value in json_data.items():
            if isinstance(value, dict):  # Handle nested dictionaries
                html += f"<tr><td>{key}</td>"  # Add the main key as a row
                # Ensure two cells for nested dictionaries
                sub_values = list(value.values())
                if len(sub_values) == 2:  # Expected two keys: 'Required' and 'Candidate'
                    html += f"<td>{sub_values[0]}</td><td>{sub_values[1]}</td></tr>"
                else:
                    raise ValueError(f"Nested dictionary for '{key}' must have exactly two keys.")
            else:  # Handle simple key-value pairs
                html += f"""
                <tr>
                    <td>{key}</td>
                    <td colspan="2">{value}</td>
                </tr>
                """

        # Close the HTML table
        html += "</table>"
        return html
    except Exception as e:
        raise ValueError(f"Error generating HTML table: {str(e)}")


def main():
    input_json = """
    {
        "Job Title": "Email Campaign Developer",
        "Candidate Name": "Anuhya Siliveru",
        "Total Years of Experience": {
            "Required": "3 - 6 Years",
            "Candidate": "9+ years"
        },
        "Key Skills": {
            "Required": "Digital Marketing, Emarsys, excel, SQL",
            "Candidate": "Spark, AWS, Snowflake, JIRA, Confluence, Python, Hive, Scala, Java, Struts, JavaScript, Ajax"
        },
        "Education": {
            "Required": "Not Mentioned",
            "Candidate": "Master of Science in Computers"
        },
        "Certifications": {
            "Required": "Not Mentioned",
            "Candidate": "Not Mentioned"
        },
        "Location": {
            "Job": "Not Mentioned",
            "Candidate": "Not Mentioned"
        },
        "Remote Eligible": {
            "Required": "Not Mentioned",
            "Candidate": "Not Mentioned"
        },
        "Salary": {
            "Range": "Not Mentioned",
            "Expected": "Not Mentioned"
        },
        "Responsibilities Match": "Candidate has extensive experience in software development and data management, but no specific experience in email campaign development.",
        "Language Proficiency": {
            "Required": "Not Mentioned",
            "Candidate": "Not Mentioned"
        },
        "Cultural Fit Indicators": "Not Mentioned",
        "Availability": {
            "Required": "Not Mentioned",
            "Candidate": "Not Mentioned"
        },
        "Resume Gaps": "None",
        "Overall Match Percentage": "50%"
    }
    """
    json_dict = json.loads(input_json)
    html = parse_nested_json_to_html_table(json_dict)
    print(html)


if __name__ == "__main__":
    main()